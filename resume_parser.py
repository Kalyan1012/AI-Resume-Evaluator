import os
import json
import time
from pathlib import Path
import streamlit as st
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document

# 1. Page Configuration
st.set_page_config(
    page_title="AI Resume Evaluator",
    page_icon="📄",
    layout="wide"
)

# 2. Setup Environment & Groq Client
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")

if not api_key:
    st.error("⚠️ GROQ_API_KEY is missing! Please set it in your .env file.")
    st.stop()

client = Groq(api_key=api_key)
MODEL = "openai/gpt-oss-120b"

# 3. Pydantic Schemas
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

class MatchResult(BaseModel):
    score: float
    candidate_name: str | None = None
    matching_skills: list[str] = []
    missing_important_skills: list[str] = []
    experience_requirement_met: bool = False
    final_verdict: str = ""

# 4. Helper Functions for Document Reading & LLM Parsing
def read_pdf(file_bytes):
    reader = PdfReader(file_bytes)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_bytes):
    doc = Document(file_bytes)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text

def extract_job_details(job_description_text):
    schema = JobD.model_json_schema()
    sys_prompt = f"""
    You are an HR assistant. Extract structured information from the job description.
    Return ONLY valid JSON matching this schema:
    {schema}
    Do NOT include schema wrapper keys like "properties" or "title".
    """
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": job_description_text}
        ],
        response_format={"type": "json_object"}
    )
    return JobD(**json.loads(response.choices[0].message.content))

def parse_resume(resume_text):
    schema = Resume.model_json_schema()
    sys_prompt = f"""
    You are an expert resume parser. Extract information based on meaning across sections.
    Return ONLY valid JSON matching this schema:
    {schema}
    Do not invent facts. Return empty lists/nulls if missing.
    """
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": resume_text}
        ],
        response_format={"type": "json_object"}
    )
    return Resume(**json.loads(response.choices[0].message.content))

def evaluate_match(job, resume):
    schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter. Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}

    Return ONLY valid JSON matching this schema:
    {schema}

    Ensure keys are populated: candidate_name, matching_skills, missing_important_skills, experience_requirement_met, score (0-100), and final_verdict.
    """
    response = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return MatchResult(**json.loads(response.choices[0].message.content))

# 5. UI Layout
st.title("🎯 AI Resume Evaluator & Candidate Ranker")
st.markdown("Upload resumes against a job description to extract metrics and candidate fit scores.")

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. Job Description")
    default_jd = """Software Development Engineer (SDE-I) at Amazon.
Requires experience with Java, Python, or C++, Data Structures, Algorithms, OOP principles, and a Bachelor's degree in Computer Science/STEM.
Preferred: AWS Cloud, AI productivity tools, SQL/NoSQL databases, and Git."""
    
    jd_input = st.text_area("Paste Job Description here", value=default_jd, height=280)

with col2:
    st.subheader("2. Candidate Resumes")
    uploaded_files = st.file_uploader(
        "Upload PDF or DOCX Resumes",
        type=["pdf", "docx"],
        accept_multiple_files=True
    )

st.divider()

if st.button("🚀 Evaluate Resumes", type="primary", use_container_width=True):
    if not jd_input.strip():
        st.warning("Please provide a Job Description.")
    elif not uploaded_files:
        st.warning("Please upload at least one PDF or DOCX resume.")
    else:
        with st.status("Extracting requirements and evaluating candidates...", expanded=True) as status:
            st.write("🔍 Extracting Job Description requirements...")
            parsed_jd = extract_job_details(jd_input)
            
            results = []
            
            for index, uploaded_file in enumerate(uploaded_files):
                st.write(f"📄 Processing resume ({index+1}/{len(uploaded_files)}): **{uploaded_file.name}**")
                
                # Read Text
                if uploaded_file.name.endswith(".pdf"):
                    resume_text = read_pdf(uploaded_file)
                else:
                    resume_text = read_docx(uploaded_file)
                
                # LLM Parsing & Evaluation
                parsed_res = parse_resume(resume_text)
                time.sleep(2)  # Short delay to manage rate limits
                
                match_res = evaluate_match(parsed_jd, parsed_res)
                time.sleep(2)
                
                results.append({
                    "file_name": uploaded_file.name,
                    "candidate_name": match_res.candidate_name or parsed_res.name or uploaded_file.name,
                    "score": match_res.score,
                    "matching_skills": match_res.matching_skills,
                    "missing_skills": match_res.missing_important_skills,
                    "exp_met": match_res.experience_requirement_met,
                    "verdict": match_res.final_verdict
                })
            
            status.update(label="Evaluation Complete!", state="complete", expanded=False)

        # Sort Candidates by Score
        results.sort(key=lambda x: x["score"], reverse=True)

        st.subheader("📊 Candidate Ranking & Score Summary")
        
        # Display Cards for Ranked Candidates
        for rank, c in enumerate(results, start=1):
            with st.container(border=True):
                c1, c2 = st.columns([1, 3])
                
                with c1:
                    st.metric(label=f"Rank #{rank} Fit Score", value=f"{c['score']}%")
                    st.caption(f"📁 `{c['file_name']}`")
                
                with c2:
                    st.subheader(c["candidate_name"])
                    st.markdown(f"**Verdict:** {c['verdict']}")
                    
                    m_col1, m_col2 = st.columns(2)
                    with m_col1:
                        st.markdown("**✅ Matching Skills:**")
                        if c["matching_skills"]:
                            st.write(", ".join(c["matching_skills"]))
                        else:
                            st.write("None identified")
                            
                    with m_col2:
                        st.markdown("**❌ Missing Skills:**")
                        if c["missing_skills"]:
                            st.write(", ".join(c["missing_skills"]))
                        else:
                            st.write("None identified")